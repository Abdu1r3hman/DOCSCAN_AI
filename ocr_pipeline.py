import os
import sys
import argparse
import re
import cv2
import numpy as np

os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
os.environ["PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK"] = "True"
os.environ["FLAGS_use_onednn"] = "0"
os.environ["FLAGS_use_mkldnn"] = "0"
os.environ["PADDLE_PIR_MODE"] = "0"
import logging
logging.disable(logging.DEBUG) 

import pytesseract
from pytesseract import Output
from PIL import Image
import docx
from docx.shared import Pt, Inches
from paddleocr import PaddleOCR
from spellchecker import SpellChecker

tesseract_cmd_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
if os.path.exists(tesseract_cmd_path):
    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd_path

# Global Spellchecker instance (Pre-loaded for high performance)
spell = SpellChecker()
spell.word_frequency.load_words(['maglev', 'homogenous', 'heterogeneous', 'catalyst', 'frictional', 'gravitational', 'paddlex'])

def smart_autocorrect(text):
    """
    A conservative spell checker that fixes common OCR typos (e.g. 'mouing' -> 'moving')
    but ignores numbers, acronyms, and preserves Word-Case formatting. 
    It explicitly protects known domain terms.
    """
    if not text: return ""
    
    words = text.split(" ")
    corrected = []
    
    for word in words:
        if not word: 
            corrected.append("")
            continue
            
        clean_word = re.sub(r'[^\w\s]', '', word)
        
        # Don't touch numbers, acronyms, short words, or URLs
        if clean_word.isupper() or any(c.isdigit() for c in clean_word) or len(clean_word) < 4:
            corrected.append(word)
            continue
            
        if spell.unknown([clean_word.lower()]):
            candidate = spell.correction(clean_word.lower())
            if candidate and abs(len(clean_word) - len(candidate)) <= 2:
                if word[0].isupper():
                    candidate = candidate.capitalize()
                
                # Reconstruct with original punctuation
                new_word = word.replace(clean_word, candidate)
                corrected.append(new_word)
            else:
                corrected.append(word)
        else:
            corrected.append(word)
            
    return " ".join(corrected)

def is_bullet_point(text):
    # Detects lines starting with bullets, arrows, or dashes
    clean = text.strip()
    return re.match(r'^([\-\*•→>]|(\d+\.?\s+))', clean)

def is_semantic_heading(text):
    clean_text = text.strip()
    text_lower = clean_text.lower()
    
    # Matches Q1, Q2, Qa, Qb, Q., Ans:, Topic:, Similarity, Difference, Week-X
    # Added 'Field Notes', 'Summary', and generic title patterns
    if re.match(r'^(q\s*[\d\w\.]|ans\b|topic\b|similarity|difference|week-?\d|field\s*notes|summary|observation)', text_lower):
        return True
    if clean_text.endswith(':-') or clean_text.endswith(':') or clean_text.endswith('-'):
        # Catch headers like 'Internal:-' or 'Field Notes-'
        if len(clean_text) < 40:
             return True
    if clean_text.isupper() and len(clean_text) > 3 and len(clean_text.split()) < 5:
        return True
    return False

def enhance_image_for_ocr(image_path):
    print("[INFO] Applying mild OpenCV preprocessing to enhance faint text...")
    img = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE) 
    
    # Upscale by 1.5x to improve faint stroke reading
    img = cv2.resize(img, None, fx=1.5, fy=1.5, interpolation=cv2.INTER_CUBIC)
    
    # Apply CLAHE (Contrast Limited Adaptive Histogram Equalization)
    clahe = cv2.createCLAHE(clipLimit=1.5, tileGridSize=(8, 8))
    img = clahe.apply(img)
    
    # Slight blur to kill the salt-and-pepper artifacts that CLAHE brings up in shadows
    img = cv2.GaussianBlur(img, (3, 3), 0)
    
    temp_path = os.path.splitext(image_path)[0] + "_enhanced_temp.png"
    cv2.imwrite(temp_path, img)
    return temp_path

_ocr_instances = {}

def get_paddle_ocr(lang='en'):
    global _ocr_instances
    if lang not in _ocr_instances:
        print(f"[SYSTEM] Initializing PaddleOCR for language: {lang}...")
        # Disable connectivity checks for cloud stability
        os.environ["PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK"] = "True"
        # Force Lightweight Mobile Models to save RAM (Fixes 16Gi Memory Limit Exceeded)
        _ocr_instances[lang] = PaddleOCR(
            use_textline_orientation=False, 
            lang=lang, 
            enable_mkldnn=False,
            ocr_version='PP-OCRv4'
        )
    return _ocr_instances[lang]

def process_handwritten_document(image_path, doc, language='en', ocr_instance=None):
    print(f"[INFO] Processing Handwritten Document with PaddleOCR (Lang: {language})...")
    enhanced_path = enhance_image_for_ocr(image_path)
    
    ocr = ocr_instance if ocr_instance else get_paddle_ocr(language)
                    
    try:
        # Standard PaddleOCR method (removed cls=True for universal compatibility)
        results = ocr.ocr(enhanced_path)
    except (AttributeError, TypeError):
        # Fallback for older/specific versions
        results = ocr.predict(enhanced_path)
    
    if not results:
        print("[WARNING] PaddleOCR returned empty results!")
        return
        
    # Standardize result format (PaddleOCR sometimes returns a list of lists)
    if isinstance(results, list) and len(results) > 0 and isinstance(results[0], list):
        results = results[0]
    
    parsed_data = []
    
    for res in results:
        try:
            if hasattr(res, 'keys') and 'rec_texts' in res:
                texts = res['rec_texts']
                boxes = res['rec_boxes'] if 'rec_boxes' in res else res['dt_polys']
            elif 'rec_texts' in dict(res):
                texts = dict(res)['rec_texts']
                boxes = dict(res)['rec_boxes'] if 'rec_boxes' in dict(res) else dict(res)['dt_polys']
            else:
                texts = res['rec_texts']
                boxes = res['rec_boxes'] if 'rec_boxes' in res else res['dt_polys']

            for text, box in zip(texts, boxes):
                if isinstance(text, str):
                    if len(text.strip()) <= 2 and text.lower() not in ['a', 'i', 'an', 'no', 'of', 'to', 'in', 'is', 'it', 'on', 'be', 'by', 'my', 'me', 'do', 'go', 'we']:
                        if not re.match(r'^[Q\d]', text):
                            continue
                            
                    if len(box) == 4 and not hasattr(box[0], '__iter__'):
                        top_left_x = float(box[0])
                        top_left_y = float(box[1])
                        width = float(box[2]) - float(box[0])
                        height = float(box[3]) - float(box[1])
                    else:
                        top_left_x = float(box[0][0])
                        top_left_y = float(box[0][1])
                        width = float(box[1][0]) - float(box[0][0])
                        height = float(box[3][1]) - float(box[0][1])
                        
                    parsed_data.append({"text": text, "x": top_left_x, "y": top_left_y, "width": width, "height": height})
        except Exception as e:
            pass

    if not parsed_data:
        print("[WARNING] No handwritten text was found!")
        return

    parsed_data.sort(key=lambda item: item['y'])
    lines = []
    current_line = []
    current_y = parsed_data[0]['y']
    
    for item in parsed_data:
        if abs(item['y'] - current_y) < 20: 
            current_line.append(item)
        else:
            current_line.sort(key=lambda i: i['x'])
            lines.append(current_line)
            current_line = [item]
            current_y = item['y']
            
    if current_line:
        current_line.sort(key=lambda i: i['x'])
        lines.append(current_line)
        
    all_xs = [item['x'] for item in parsed_data]
    all_heights = [item['height'] for item in parsed_data]
    
    page_min_x = min(all_xs) if all_xs else 0
    page_max_x = max([item['x'] + item['width'] for item in parsed_data]) if parsed_data else 1000
    page_width = page_max_x - page_min_x
    if page_width == 0: page_width = 1000
    
    average_height = sum(all_heights) / len(all_heights) if all_heights else 10

    vertical_gaps = []
    for i in range(1, len(lines)):
        gap = lines[i][0]['y'] - lines[i-1][0]['y']
        if gap > 0: vertical_gaps.append(gap)
    avg_gap = sum(vertical_gaps) / len(vertical_gaps) if vertical_gaps else 30

    p = None
    for i, line_items in enumerate(lines):
        # Spatially aware line joining (preserves tables and 2-column layouts via Tabs)
        line_text = ""
        last_x_end = None
        has_large_gap = False
        
        for k, item in enumerate(line_items):
            if k == 0:
                line_text += item['text']
            else:
                pixel_gap = item['x'] - last_x_end
                
                # Single tab injection for column separations
                if pixel_gap > (page_width * 0.04):
                    line_text += "\t"
                    has_large_gap = True
                else:
                    line_text += " "
                    
                line_text += item['text']
                
            last_x_end = item['x'] + item['width']
        first_word_x = line_items[0]['x']
        relative_ratio = (first_word_x - page_min_x) / page_width
        is_margin_bullet = relative_ratio < 0.05
        
        line_height = max([item['height'] for item in line_items])
        current_y = line_items[0]['y']
        
        # Apply intelligent spelling correction before heading analysis and writing
        line_text = smart_autocorrect(line_text)
        
        is_large_text = line_height > (average_height * 1.6)
        
        is_bullet = is_bullet_point(line_text)
        # Heading triggers: Semantic Match OR physically large text
        is_heading = is_semantic_heading(line_text) or is_large_text
        
        is_new_paragraph = True
        
        if p is not None and i > 0:
            prev_y = lines[i-1][0]['y']
            y_diff = current_y - prev_y
            
            # Bullets always start new paragraphs
            if is_bullet:
                is_new_paragraph = True
            # If the gap is relatively normal and we aren't explicitly creating a heading or a Table Row
            # or a margin-started line (which usually indicates a list/bullet/new logical block)
            elif y_diff < (avg_gap * 1.4) and not is_heading and not has_large_gap and not is_margin_bullet:
                is_new_paragraph = False
        
        if is_new_paragraph:
            p = doc.add_paragraph()
            
            # Ethnographic readability: Add more space before headings
            if is_heading:
                p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            # Set the physical indentation
            first_word_x = line_items[0]['x']
            relative_ratio = (first_word_x - page_min_x) / page_width
            
            # Bullet/List logic: Apply a hanging indent so text wraps cleanly
            base_indent = min(relative_ratio * 6.0, 3.5)
            if is_bullet:
                # Add extra padding for the bullet symbol
                p.paragraph_format.left_indent = Inches(base_indent + 0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
            elif base_indent > 0.15:
                p.paragraph_format.left_indent = Inches(base_indent)
                
            p.alignment = 0 # Left align
            
            # 5-Column Safety Tab Stops
            tab_stops = p.paragraph_format.tab_stops
            for stop_idx in range(1, 6):
                tab_stops.add_tab_stop(Inches(1.2 * stop_idx))
            
            # Application of Bolding Rules
            if is_large_text:
                run = p.add_run(line_text)
                run.bold = True
            elif is_heading:
                # Detected a tag like 'Q1:', 'Ans:', or ends with ':-'
                if ":" in line_text or line_text.endswith("-"):
                    # Catch 'Internal:-' or 'Field Notes-'
                    split_char = ":" if ":" in line_text else "-"
                    parts = line_text.split(split_char, 1)
                    tag_run = p.add_run(parts[0] + split_char)
                    tag_run.bold = True
                    if len(parts) > 1:
                        p.add_run(parts[1])
                elif len(line_text) < 40:
                    run = p.add_run(line_text)
                    run.bold = True
                    run.italic = True
                else:
                    p.add_run(line_text)
            else:
                p.add_run(line_text)
        else:
            # Append space and continue on the same paragraph block!
            p.add_run(" " + line_text)

    # Cleanup temp file
    if os.path.exists(enhanced_path):
        os.remove(enhanced_path)

def process_printed_document(image_path, doc):
    print("[INFO] Processing Printed Document with Tesseract...")
    img = Image.open(image_path)
    img_width, img_height = img.size
    
    data = pytesseract.image_to_data(img, output_type=Output.DICT)
    n_boxes = len(data['level'])
    
    current_block = -1
    current_par = -1
    current_line = -1
    
    line_text = []
    line_left = []
    
    for i in range(n_boxes):
        word = data['text'][i].strip()
        conf = int(data['conf'][i])
        
        # Tesseract confidence threshold to skip noise
        if word and conf > 40:
            block_num = data['block_num'][i]
            par_num = data['par_num'][i]
            line_num = data['line_num'][i]
            left = data['left'][i]
            
            if line_num != current_line or par_num != current_par or block_num != current_block:
                if line_text:
                    joined_text = " ".join(line_text)
                    min_left = min(line_left) if line_left else 0
                    
                    p = doc.add_paragraph()
                    
                    # Layout alignment logic via Tesseract bounds
                    relative_ratio = min_left / img_width
                    if relative_ratio > 0.3:
                        p.alignment = 1 # Center
                    else:
                        p.alignment = 0 # Left
                        if relative_ratio > 0.05:
                            p.paragraph_format.left_indent = Inches(relative_ratio * 6.0)
                            
                    run = p.add_run(joined_text)
                    # Basic semantic heading for printed
                    if is_semantic_heading(joined_text):
                        run.bold = True
                        
                current_block = block_num
                current_par = par_num
                current_line = line_num
                line_text = [word]
                line_left = [left]
            else:
                line_text.append(word)
                line_left.append(left)
                
    if line_text:
        joined_text = " ".join(line_text)
        p = doc.add_paragraph()
        p.add_run(joined_text)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("image_path")
    parser.add_argument("--type", choices=['printed', 'handwritten', 'auto'], default='auto')
    parser.add_argument("--lang", default='en')
    args = parser.parse_args()
    
    if not os.path.exists(args.image_path):
        print(f"[ERROR] Could not find the image '{args.image_path}'.")
        return
        
    output_docx = os.path.splitext(args.image_path)[0] + "_final.docx"
    doc = docx.Document()
    
    doc_type = args.type
    if doc_type == 'auto': doc_type = 'handwritten' 
        
    if doc_type == 'printed':
        process_printed_document(args.image_path, doc)
    elif doc_type == 'handwritten':
        process_handwritten_document(args.image_path, doc, language=args.lang)
        
    try:
        doc.save(output_docx)
        print(f"\n[SUCCESS] Formatted unified workflow complete!")
        print(f"File saved to: {output_docx}")
    except PermissionError:
        print(f"\n[ERROR] Permission Denied: Could not save to '{output_docx}'.")

if __name__ == "__main__":
    main()
