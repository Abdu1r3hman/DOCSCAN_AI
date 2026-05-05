import os
import sys
import argparse
import base64
import re
import json
import urllib.request
import docx
from docx.shared import Pt, Inches

try:
    import google.generativeai as genai
    from dotenv import load_dotenv
    load_dotenv()
    HAS_GEMINI = True
except ImportError:
    HAS_GEMINI = False

def parse_markdown_to_docx(markdown_text, doc):
    """
    Renders LLM/VLM markdown text natively into MS Word formatted tables and paragraphs.
    """
    lines = markdown_text.strip().split('\n')
    
    in_table = False
    table_data = []
    
    for line in lines:
        line = line.strip()
        if not line:
            if not in_table:
                continue
            
        # Table Detection
        if line.startswith('|') and line.endswith('|'):
            in_table = True
            if '---' in line:
                continue
            
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
            continue
        
        # When Table finishes
        if in_table and not line.startswith('|'):
            in_table = False
            if table_data:
                rows = len(table_data)
                cols = len(table_data[0]) if rows > 0 else 0
                if cols > 0:
                    word_table = doc.add_table(rows=rows, cols=cols)
                    word_table.style = 'Table Grid'
                    for r in range(rows):
                        for c in range(min(cols, len(table_data[r]))):
                            word_table.rows[r].cells[c].text = table_data[r][c]
                table_data = [] 
                doc.add_paragraph() 
        
        # Heading Detection
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            continue
            
        # Basic Bold Parsing for Standard Paragraphs
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

def process_with_vlm(image_path, doc):
    print("\n[INFO] Initiating Vision-Language Model (VLM) Pipeline...")
    print("[INFO] This bypasses traditional OCR completely for superior spatial awareness.")
    
    system_prompt = """
    You are an expert Document Digitization AI. I will provide an image of a document (handwritten or printed).
    Your task is to transcribe the text exactly as it appears and format it beautifully using Markdown.
    
    Instructions:
    1. Transcribe the full text accurately. Fix obvious spelling mistakes in handwriting, but preserve scientific terms.
    2. Format the output clearly using standard Markdown.
    3. Bold all questions, headers, and topic titles (e.g. **Q1:** or **Topic:**).
    4. If you see tabular data or column comparisons, explicitly format them as true Markdown tables (e.g., | Col1 | Col2 |).
    5. DO NOT hallucinate extra information. Output ONLY the transcribed and formatted markdown layout.
    """
    
    # 1. Try Gemini Vision (Zero load on your laptop, instant, highly accurate)
    api_key = os.getenv("GEMINI_API_KEY")
    vlm_success = False
    
    if HAS_GEMINI and api_key:
        print("[INFO] Attempting lightweight extraction via Gemini 1.5 Flash Vision (0% Laptop Compute)...")
        try:
            from PIL import Image
            img = Image.open(image_path)
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-2.5-flash")
            response = model.generate_content([system_prompt, img])
            print("[INFO] VLM transcription successful via Gemini!")
            parse_markdown_to_docx(response.text, doc)
            vlm_success = True
        except Exception as e:
            print(f"[WARNING] Gemini Vision failed: {e}")
            
    # 2. Try Ollama Vision (Local fallback, requires 'llava' or 'moondream')
    if not vlm_success:
        print("[INFO] Attempting local VLM extraction via Ollama (Requires 'llava' model)...")
        try:
            with open(image_path, "rb") as f:
                img_b64 = base64.b64encode(f.read()).decode('utf-8')
                
            data = json.dumps({
                "model": "llava", 
                "prompt": system_prompt,
                "images": [img_b64],
                "stream": False
            }).encode('utf-8')
            
            req = urllib.request.Request("http://localhost:11434/api/generate", data=data, headers={'Content-Type': 'application/json'})
            with urllib.request.urlopen(req, timeout=120) as response:
                result = json.loads(response.read().decode('utf-8'))
                result_text = result.get('response', '')
                print("[INFO] VLM transcription successful via local Ollama!")
                parse_markdown_to_docx(result_text, doc)
                vlm_success = True
        except Exception as e:
            print(f"[ERROR] Local Ollama VLM failed: {e}")
            print("\n[HINT] To run locally, open terminal and type: 'ollama run llava'")
            
    if not vlm_success:
        print("\n[CRITICAL ERROR] Both VLM engines failed.")
        doc.add_paragraph("[ERROR] System could not process the image using VLM.")
        doc.add_paragraph("Please ensure you have an active internet connection for Gemini, OR you have pulled the 'llava' model via Ollama.")

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("image_path")
    args = parser.parse_args()
    
    if not os.path.exists(args.image_path):
        print(f"[ERROR] Could not find the image '{args.image_path}'.")
        return
        
    output_docx = os.path.splitext(args.image_path)[0] + "_vlm_direct.docx"
    doc = docx.Document()
    
    process_with_vlm(args.image_path, doc)
        
    try:
        doc.save(output_docx)
        print(f"\n[SUCCESS] Document created seamlessly via VLM: {output_docx}")
    except PermissionError:
        print(f"\n[ERROR] File is locked by MS Word.")

if __name__ == "__main__":
    main()
