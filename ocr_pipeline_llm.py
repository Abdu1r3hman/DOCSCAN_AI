import os
import sys
import argparse
import re
import cv2

os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
import logging
logging.disable(logging.DEBUG) 

import docx
from docx.shared import Pt, Inches
from paddleocr import PaddleOCR

# To use the LLM, we import Google Generative AI (Gemini Flash is free/cheap)
try:
    import google.generativeai as genai
    from dotenv import load_dotenv
    load_dotenv()
    HAS_LLM = True
except ImportError:
    HAS_LLM = False

def enhance_image_for_ocr(image_path):
    print("[INFO] OpenCV preprocessing for faint text...")
    img = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE) 
    img = cv2.resize(img, None, fx=1.5, fy=1.5, interpolation=cv2.INTER_CUBIC)
    clahe = cv2.createCLAHE(clipLimit=1.5, tileGridSize=(8, 8))
    img = clahe.apply(img)
    img = cv2.GaussianBlur(img, (3, 3), 0)
    
    temp = os.path.splitext(image_path)[0] + "_enhanced_temp.png"
    cv2.imwrite(temp, img)
    return temp

def parse_markdown_to_docx(markdown_text, doc):
    """
    Renders LLM markdown text natively into MS Word formatted tables and paragraphs.
    """
    lines = markdown_text.strip().split('\n')
    
    in_table = False
    table_data = []
    
    for line in lines:
        line = line.strip()
        if not line:
            # Empty line, reset spaces if not in table
            if not in_table:
                continue
            
        # Table Detection
        if line.startswith('|') and line.endswith('|'):
            in_table = True
            # Ignore markdown table separator block (e.g. |---|---|)
            if '---' in line:
                continue
            
            # Split cells
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
            continue
        
        # When Table finishes
        if in_table and not line.startswith('|'):
            in_table = False
            if table_data:
                rows = len(table_data)
                cols = len(table_data[0]) if rows > 0 else 0
                if cols > 0:
                    word_table = doc.add_table(rows=rows, cols=cols)
                    word_table.style = 'Table Grid'
                    for r in range(rows):
                        for c in range(min(cols, len(table_data[r]))):
                            word_table.rows[r].cells[c].text = table_data[r][c]
                table_data = [] # Reset
                doc.add_paragraph() # Add gap after table
        
        # Heading Detection
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            continue
            
        # Basic Bold Parsing for Standard Paragraphs
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

def process_handwritten_document(image_path, doc):
    print("[INFO] Extracting Raw Text Coordinates with PaddleOCR...")
    enhanced_path = enhance_image_for_ocr(image_path)
    ocr = PaddleOCR(use_textline_orientation=True, lang='en')
    results = ocr.predict(enhanced_path)
    
    # We will build a Raw Messy String maintaining vertical order for the LLM
    parsed_data = []
    for res in results:
        try:
            if hasattr(res, 'keys') and 'rec_texts' in res:
                texts = res['rec_texts']
                boxes = res['rec_boxes'] if 'rec_boxes' in res else res['dt_polys']
            elif 'rec_texts' in dict(res):
                texts = dict(res)['rec_texts']
                boxes = dict(res)['rec_boxes'] if 'rec_boxes' in dict(res) else dict(res)['dt_polys']
            else:
                texts = res['rec_texts']
                boxes = res['rec_boxes'] if 'rec_boxes' in res else res['dt_polys']

            for text, box in zip(texts, boxes):
                if isinstance(text, str):
                    if len(box) == 4 and not hasattr(box[0], '__iter__'):
                        top_left_y = float(box[1])
                    else:
                        top_left_y = float(box[0][1])
                    parsed_data.append({"text": text, "y": top_left_y})
        except Exception:
            pass

    if not parsed_data:
        print("[WARNING] No text found!")
        return
        
    parsed_data.sort(key=lambda item: item['y'])
    raw_ocr_string = "\n".join([item['text'] for item in parsed_data])
    
    if os.path.exists(enhanced_path):
        os.remove(enhanced_path)

    # ================= LLM INTEGRATION =================
    print("\n[INFO] Attempting to use local Ollama LLM for cognitive formatting...")
    
    system_prompt = f"""
    You are a highly advanced Note Formatting assistant. I will provide messy OCR raw text from handwritten study notes.
    Your instructions:
    1. Contextual Spelling: Fix OCR typos directly in the text (e.g. 'mouing' -> 'moving', 'Da maglev' -> 'Q2: Maglev'.) Do not break scientific words.
    2. Format clearly: Use markdown.
    3. Bold questions and question numbers using **Q1:** or **Topic:** format.
    4. Detect Column comparisons (e.g., Homogenous vs Heterogenous features) and explicitly format them as true Markdown tables (e.g., | Col1 | Col2 |).
    5. DO NOT hallucinate extra information. Output strictly the formatted markdown layout.

    RAW TEXT:
    {raw_ocr_string}
    """
    
    ollama_success = False
    try:
        import json
        import urllib.request
        
        data = json.dumps({
            "model": "llama3",
            "prompt": system_prompt,
            "stream": False
        }).encode('utf-8')
        
        req = urllib.request.Request("http://localhost:11434/api/generate", data=data, headers={'Content-Type': 'application/json'})
        with urllib.request.urlopen(req, timeout=120) as response:
            result = json.loads(response.read().decode('utf-8'))
            result_text = result.get('response', '')
            print("\n[INFO] AI formatting generated successfully via local Ollama!")
            parse_markdown_to_docx(result_text, doc)
            ollama_success = True
    except Exception as e:
        print(f"\n[WARNING] Local Ollama failed or not running: {e}")
        
    if not ollama_success:
        api_key = os.getenv("GEMINI_API_KEY")
        if HAS_LLM and api_key:
            print("\n[INFO] Sending OCR string to Gemini Flash LLM for cognitive formatting...")
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-2.5-flash")
            response = model.generate_content(system_prompt)
            print("\n[INFO] AI formatting generated successfully via Gemini!")
            parse_markdown_to_docx(response.text, doc)
        else:
            print("\n[ERROR] No LLM Configured! Please ensure Ollama is running or Gemini API key is set.")
            print("[RAW OCR FALLBACK]")
            doc.add_paragraph(raw_ocr_string)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("image_path")
    args = parser.parse_args()
    
    if not os.path.exists(args.image_path):
        return
        
    output_docx = os.path.splitext(args.image_path)[0] + "_smart_llm.docx"
    doc = docx.Document()
    
    process_handwritten_document(args.image_path, doc)
        
    try:
        doc.save(output_docx)
        print(f"\n[SUCCESS] Saved to Document: {output_docx}")
    except PermissionError:
        print(f"\n[ERROR] File is locked by MS Word.")

if __name__ == "__main__":
    main()
