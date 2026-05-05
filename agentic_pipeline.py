import os
import sys

# Ensure stdout handles emojis on Windows
if hasattr(sys.stdout, 'reconfigure') and sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

import argparse
import json
import re
from PIL import Image
import docx
from docx.shared import Pt, Inches

try:
    from openai import OpenAI
    from dotenv import load_dotenv
    load_dotenv(override=True)
    HAS_OPENROUTER = True
except ImportError:
    HAS_OPENROUTER = False

# ==============================================================================
# PHASE 2: PURELY AGENTIC SYSTEM ARCHITECTURE
# ==============================================================================
# This system implements an autonomous Multi-Agent Workflow using a Supervisor
# Pattern. The agents collaborate, critique each other, and iterate until the
# output meets quality standards.
# ==============================================================================

class BaseAgent:
    def __init__(self, name, role, model="google/gemini-2.5-flash", temperature=0.3):
        self.name = name
        self.role = role
        self.model_name = model
        self.temperature = temperature
        
        if not HAS_OPENROUTER or not os.getenv("OPENROUTER_API_KEY"):
            raise EnvironmentError("Agentic System requires openai and an OPENROUTER_API_KEY in .env")
            
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=os.getenv("OPENROUTER_API_KEY"),
        )

    def act(self, prompt, image=None):
        print(f"  🤖 [{self.name}] Analyzing task...")
        messages = [
            {"role": "system", "content": self.role}
        ]
        
        if image:
            import io, base64
            buffered = io.BytesIO()
            if image.mode == 'RGBA':
                image = image.convert('RGB')
            image.save(buffered, format="JPEG")
            img_b64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
            
            content = [
                {"type": "text", "text": prompt},
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/jpeg;base64,{img_b64}",
                    },
                },
            ]
        else:
            content = prompt

        messages.append({"role": "user", "content": content})
        
        chat_completion = self.client.chat.completions.create(
            messages=messages,
            model=self.model_name,
            temperature=self.temperature,
            max_tokens=4000
        )
        return chat_completion.choices[0].message.content.strip()

def parse_markdown_to_docx(markdown_text, doc):
    """
    Renders LLM/VLM markdown text natively into MS Word formatted tables and paragraphs.
    """
    lines = markdown_text.strip().split('\n')
    in_table = False
    table_data = []
    
    for line in lines:
        line = line.strip()
        if not line:
            if not in_table:
                continue
            
        if line.startswith('|') and line.endswith('|'):
            in_table = True
            if '---' in line:
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
            continue
        
        if in_table and not line.startswith('|'):
            in_table = False
            if table_data:
                rows = len(table_data)
                cols = len(table_data[0]) if rows > 0 else 0
                if cols > 0:
                    word_table = doc.add_table(rows=rows, cols=cols)
                    word_table.style = 'Table Grid'
                    for r in range(rows):
                        for c in range(min(cols, len(table_data[r]))):
                            word_table.rows[r].cells[c].text = table_data[r][c]
                table_data = [] 
                doc.add_paragraph() 
        
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            continue
            
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

    if in_table and table_data:
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0
        if cols > 0:
            word_table = doc.add_table(rows=rows, cols=cols)
            word_table.style = 'Table Grid'
            for r in range(rows):
                for c in range(min(cols, len(table_data[r]))):
                    word_table.rows[r].cells[c].text = table_data[r][c]

# ==============================================================================
# AGENTIC WORKFLOW EXECUTION
# ==============================================================================

def run_agentic_mas(image_path, output_docx):
    print("\n[SYSTEM] Initializing Purely Agentic Multi-Agent System (MAS)...")
    print("[SYSTEM] Architecture: VisionAgent -> QAAgent -> FormatterAgent\n")
    
    try:
        img = Image.open(image_path)
    except Exception as e:
        print(f"[ERROR] Could not open image: {e}")
        return

    # 1. Initialize Agents
    vision_agent = BaseAgent(
        name="VisionExtractor", 
        role="You are a highly perceptive Vision Extractor Agent. Your ONLY job is to transcribe every single word and structural element (tables, lists) from the image. Do not format it beautifully, just ensure complete data capture. Correct obvious spelling mistakes if it is handwriting.",
        temperature=0.1
    )
    
    qa_agent = BaseAgent(
        name="QAReviewer",
        role="""You are a strict Quality Assurance Agent. Review the extracted text.
        Check for logical inconsistencies, broken sentences, or missing structures.
        You MUST respond strictly with a valid JSON block containing two keys:
        {
            "status": "PASS" or "FAIL",
            "feedback": "Explain what is missing or incorrect, or write 'Looks perfect' if PASS"
        }
        Do not include markdown blocks like ```json in your response, just the raw JSON text.""",
        temperature=0.2
    )
    
    formatter_agent = BaseAgent(
        name="Formatter",
        role="You are a Master Formatter Agent. Take the raw QA-approved text and format it into beautiful, standardized Markdown. Use headers (#, ##), bold tags (**text**), and explicitly format tabular data into Markdown tables (| col | col |). Do NOT hallucinate extra information. Output ONLY the markdown.",
        temperature=0.1
    )

    # 2. Autonomous Loop (Max 2 Iterations to prevent infinite loops)
    raw_text = ""
    max_retries = 2
    
    for iteration in range(1, max_retries + 1):
        print(f"\n--- [AUTONOMOUS LOOP: ITERATION {iteration}] ---")
        
        prompt = "Extract all text and tabular data from this image."
        if iteration > 1:
            prompt += f"\n\nCRITICAL FEEDBACK FROM QA AGENT TO FIX:\n{qa_feedback}"
            
        # Agent 1 Acts
        raw_text = vision_agent.act(prompt, image=img)
        print(f"  ✅ [VisionExtractor] Extracted {len(raw_text)} characters of raw data.")
        
        # Agent 2 Critiques
        qa_eval_string = qa_agent.act(f"Review this extracted text for quality:\n\n{raw_text}")
        
        try:
            # Clean JSON response from LLM quirks
            clean_json = qa_eval_string.replace('```json', '').replace('```', '').strip()
            qa_json = json.loads(clean_json)
            status = qa_json.get("status", "PASS").upper()
            qa_feedback = qa_json.get("feedback", "No feedback provided.")
            
            print(f"  🔍 [QAReviewer] Audit Status: {status}")
            print(f"  💬 [QAReviewer] Feedback: {qa_feedback}")
            
            if status == "PASS":
                print("  ✅ [SYSTEM] QA Approved! Breaking loop.")
                break
            else:
                print("  🔄 [SYSTEM] QA Failed. Sending feedback back to VisionExtractor...")
                
        except json.JSONDecodeError:
            print("  ⚠️ [QAReviewer] Failed to output valid JSON. Assuming PASS to avoid infinite loop.")
            break
            
    # 3. Final Formatting
    print("\n--- [FINALIZATION PHASE] ---")
    final_markdown = formatter_agent.act(f"Format this approved text into perfect Markdown:\n\n{raw_text}")
    print(f"  ✅ [Formatter] Markdown generation complete.")
    
    print("\n=== FINAL_OUTPUT_START ===")
    print(final_markdown)
    print("=== FINAL_OUTPUT_END ===\n")
    
    # 4. Save to Disk
    print("\n[SYSTEM] Writing output to Word Document...")
    doc = docx.Document()
    parse_markdown_to_docx(final_markdown, doc)
    
    try:
        doc.save(output_docx)
        print(f"[SUCCESS] Agentic Workflow Complete! File saved: {output_docx}")
    except PermissionError:
        print(f"\n[ERROR] File is locked by MS Word.")

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("image_path")
    args = parser.parse_args()
    
    if not os.path.exists(args.image_path):
        print(f"[ERROR] Could not find image '{args.image_path}'.")
        return
        
    output_docx = os.path.splitext(args.image_path)[0] + "_agentic_output.docx"
    
    try:
        run_agentic_mas(args.image_path, output_docx)
    except Exception as e:
        print(f"\n[CRITICAL SYSTEM ERROR] {e}")

if __name__ == "__main__":
    main()
